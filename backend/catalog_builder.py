"""
Catalog Builder for AI Library
Scans directories and builds searchable file catalog
"""

import os
import sqlite3
import hashlib
import mimetypes
from datetime import datetime
from pathlib import Path
from typing import Callable, Optional, Dict, Any

# File content extractors
import magic
from PIL import Image
from PIL.ExifTags import TAGS
import PyPDF2
from docx import Document
import openpyxl
from mutagen import File as MutagenFile

class CatalogBuilder:
    def __init__(self, directory_path: str, ai_client):
        self.directory_path = Path(directory_path)
        self.ai_client = ai_client
        self.db_path = Path.home() / '.ai_library' / 'catalog.db'
        self.supported_text_types = {
            '.txt', '.md', '.py', '.js', '.html', '.css', '.json', '.xml', '.csv',
            '.log', '.ini', '.cfg', '.conf', '.yaml', '.yml', '.sql', '.sh', '.bat'
        }
        
        # Initialize database
        self._init_database()
    
    def _init_database(self):
        """Initialize SQLite database for file catalog"""
        self.db_path.parent.mkdir(exist_ok=True)
        
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            # Create files table
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS files (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    path TEXT UNIQUE NOT NULL,
                    name TEXT NOT NULL,
                    size INTEGER,
                    modified_time TEXT,
                    file_type TEXT,
                    mime_type TEXT,
                    content_hash TEXT,
                    description TEXT,
                    content_text TEXT,
                    metadata_json TEXT,
                    indexed_time TEXT
                )
            ''')
            
            # Create search index
            cursor.execute('''
                CREATE VIRTUAL TABLE IF NOT EXISTS files_fts USING fts5(
                    name, description, content_text, content='files', content_rowid='id'
                )
            ''')
            
            conn.commit()
    
    def build_catalog(self, progress_callback: Optional[Callable] = None):
        """Build the complete file catalog"""
        print(f"🔍 Starting catalog build for: {self.directory_path}")
        
        # Get all files
        all_files = list(self._get_all_files())
        total_files = len(all_files)
        
        print(f"📊 Found {total_files} files to process")
        
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            for i, file_path in enumerate(all_files):
                try:
                    if progress_callback:
                        progress_callback(str(file_path), i, total_files)
                    
                    # Check if file already exists and hasn't changed
                    if self._is_file_already_indexed(cursor, file_path):
                        continue
                    
                    # Process the file
                    file_info = self._process_file(file_path)
                    
                    if file_info:
                        # Insert or update file record
                        cursor.execute('''
                            INSERT OR REPLACE INTO files 
                            (path, name, size, modified_time, file_type, mime_type, 
                             content_hash, description, content_text, metadata_json, indexed_time)
                            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                        ''', (
                            str(file_path),
                            file_info['name'],
                            file_info['size'],
                            file_info['modified_time'],
                            file_info['file_type'],
                            file_info['mime_type'],
                            file_info['content_hash'],
                            file_info['description'],
                            file_info['content_text'],
                            file_info['metadata_json'],
                            datetime.now().isoformat()
                        ))
                        
                        # Update FTS index
                        cursor.execute('''
                            INSERT OR REPLACE INTO files_fts (rowid, name, description, content_text)
                            VALUES (last_insert_rowid(), ?, ?, ?)
                        ''', (
                            file_info['name'],
                            file_info['description'],
                            file_info['content_text']
                        ))
                    
                    # Commit every 100 files
                    if i % 100 == 0:
                        conn.commit()
                        
                except Exception as e:
                    print(f"❌ Error processing {file_path}: {e}")
                    continue
            
            # Final commit
            conn.commit()
        
        print("✅ Catalog build complete!")
    
    def _get_all_files(self):
        """Get all files in the directory recursively"""
        for root, dirs, files in os.walk(self.directory_path):
            # Skip hidden directories and common ignore patterns
            dirs[:] = [d for d in dirs if not d.startswith('.') and d not in {
                '__pycache__', 'node_modules', '.git', '.svn', 'venv', 'env'
            }]
            
            for file in files:
                if not file.startswith('.'):  # Skip hidden files
                    yield Path(root) / file
    
    def _is_file_already_indexed(self, cursor, file_path: Path) -> bool:
        """Check if file is already indexed and hasn't changed"""
        try:
            stat = file_path.stat()
            current_hash = self._get_file_hash(file_path)
            
            cursor.execute(
                'SELECT content_hash FROM files WHERE path = ?',
                (str(file_path),)
            )
            result = cursor.fetchone()
            
            return result and result[0] == current_hash
            
        except Exception:
            return False
    
    def _process_file(self, file_path: Path) -> Optional[Dict[str, Any]]:
        """Process a single file and extract information"""
        try:
            stat = file_path.stat()
            
            # Basic file information
            file_info = {
                'name': file_path.name,
                'size': stat.st_size,
                'modified_time': datetime.fromtimestamp(stat.st_mtime).isoformat(),
                'file_type': file_path.suffix.lower(),
                'mime_type': mimetypes.guess_type(str(file_path))[0] or 'unknown',
                'content_hash': self._get_file_hash(file_path),
                'description': '',
                'content_text': '',
                'metadata_json': '{}'
            }
            
            # Extract content and metadata based on file type
            content_text, metadata = self._extract_content_and_metadata(file_path)
            file_info['content_text'] = content_text or ''
            file_info['metadata_json'] = str(metadata) if metadata else '{}'
            
            # Generate AI description
            file_info['description'] = self.ai_client.generate_description(
                str(file_path), content_text, metadata
            )
            
            return file_info
            
        except Exception as e:
            print(f"Error processing file {file_path}: {e}")
            return None
    
    def _get_file_hash(self, file_path: Path) -> str:
        """Get MD5 hash of file for change detection"""
        try:
            hash_md5 = hashlib.md5()
            with open(file_path, "rb") as f:
                # Read in chunks to handle large files
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_md5.update(chunk)
            return hash_md5.hexdigest()
        except Exception:
            return ''
    
    def _extract_content_and_metadata(self, file_path: Path) -> tuple:
        """Extract content and metadata from file"""
        try:
            file_ext = file_path.suffix.lower()
            
            # Text files
            if file_ext in self.supported_text_types:
                return self._extract_text_content(file_path)
            
            # PDF files
            elif file_ext == '.pdf':
                return self._extract_pdf_content(file_path)
            
            # Word documents
            elif file_ext in {'.docx', '.doc'}:
                return self._extract_docx_content(file_path)
            
            # Excel files
            elif file_ext in {'.xlsx', '.xls'}:
                return self._extract_excel_content(file_path)
            
            # Images
            elif file_ext in {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff'}:
                return self._extract_image_metadata(file_path)
            
            # Audio files
            elif file_ext in {'.mp3', '.wav', '.flac', '.m4a', '.ogg'}:
                return self._extract_audio_metadata(file_path)
            
            # Video files
            elif file_ext in {'.mp4', '.avi', '.mov', '.mkv', '.wmv'}:
                return self._extract_video_metadata(file_path)
            
            else:
                # Unknown file type - just return basic info
                return '', {'file_type': file_ext}
                
        except Exception as e:
            print(f"Error extracting content from {file_path}: {e}")
            return '', {}
    
    def _extract_text_content(self, file_path: Path) -> tuple:
        """Extract content from text files"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            # Limit content size
            if len(content) > 10000:
                content = content[:10000] + "..."
            
            return content, {'file_type': 'text', 'char_count': len(content)}
            
        except Exception as e:
            return '', {'error': str(e)}
    
    def _extract_pdf_content(self, file_path: Path) -> tuple:
        """Extract content from PDF files"""
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text = ''
                
                # Extract text from first few pages
                for i, page in enumerate(reader.pages[:5]):  # Limit to first 5 pages
                    text += page.extract_text() + '\n'
                
                metadata = {
                    'file_type': 'pdf',
                    'page_count': len(reader.pages),
                    'char_count': len(text)
                }
                
                return text[:5000], metadata  # Limit text size
                
        except Exception as e:
            return '', {'error': str(e)}
    
    def _extract_docx_content(self, file_path: Path) -> tuple:
        """Extract content from Word documents"""
        try:
            doc = Document(file_path)
            text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            
            metadata = {
                'file_type': 'docx',
                'paragraph_count': len(doc.paragraphs),
                'char_count': len(text)
            }
            
            return text[:5000], metadata  # Limit text size
            
        except Exception as e:
            return '', {'error': str(e)}
    
    def _extract_excel_content(self, file_path: Path) -> tuple:
        """Extract content from Excel files"""
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            content = []
            
            for sheet_name in workbook.sheetnames[:3]:  # Limit to first 3 sheets
                sheet = workbook[sheet_name]
                sheet_content = []
                
                for row in sheet.iter_rows(max_row=20, values_only=True):  # Limit rows
                    row_text = ' '.join([str(cell) for cell in row if cell is not None])
                    if row_text.strip():
                        sheet_content.append(row_text)
                
                if sheet_content:
                    content.append(f"Sheet {sheet_name}: " + ' | '.join(sheet_content))
            
            text = '\n'.join(content)
            metadata = {
                'file_type': 'excel',
                'sheet_count': len(workbook.sheetnames),
                'char_count': len(text)
            }
            
            return text[:3000], metadata  # Limit text size
            
        except Exception as e:
            return '', {'error': str(e)}
    
    def _extract_image_metadata(self, file_path: Path) -> tuple:
        """Extract metadata from image files"""
        try:
            with Image.open(file_path) as img:
                metadata = {
                    'file_type': 'image',
                    'format': img.format,
                    'size': img.size,
                    'mode': img.mode
                }
                
                # Extract EXIF data
                if hasattr(img, '_getexif') and img._getexif():
                    exif = img._getexif()
                    for tag_id, value in exif.items():
                        tag = TAGS.get(tag_id, tag_id)
                        if isinstance(value, str) and len(value) < 100:
                            metadata[f'exif_{tag}'] = value
                
                return '', metadata
                
        except Exception as e:
            return '', {'error': str(e)}
    
    def _extract_audio_metadata(self, file_path: Path) -> tuple:
        """Extract metadata from audio files"""
        try:
            audio_file = MutagenFile(file_path)
            if audio_file:
                metadata = {
                    'file_type': 'audio',
                    'length': getattr(audio_file.info, 'length', 0),
                    'bitrate': getattr(audio_file.info, 'bitrate', 0)
                }
                
                # Extract tags
                for key, value in audio_file.items():
                    if isinstance(value, list) and value:
                        metadata[key] = str(value[0])
                    elif isinstance(value, str):
                        metadata[key] = value
                
                return '', metadata
            
            return '', {'file_type': 'audio'}
            
        except Exception as e:
            return '', {'error': str(e)}
    
    def _extract_video_metadata(self, file_path: Path) -> tuple:
        """Extract metadata from video files (basic)"""
        try:
            # For now, just return basic file info
            # Could be extended with ffmpeg or similar for detailed metadata
            stat = file_path.stat()
            metadata = {
                'file_type': 'video',
                'size_mb': round(stat.st_size / (1024 * 1024), 2)
            }
            
            return '', metadata
            
        except Exception as e:
            return '', {'error': str(e)}

